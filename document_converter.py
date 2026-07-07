# document_converter.py
import os
import logging
import asyncio
from concurrent.futures import ThreadPoolExecutor
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches
from docx.enum.shape import WD_INLINE_SHAPE_TYPE

# إعداد السجلات
logger = logging.getLogger(__name__)
executor = ThreadPoolExecutor(max_workers=4)

# حل المشكلة الأولى: جعل مسار الخط مطلقاً بالنسبة لمجلد هذا الملف البرمجي نفسه
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FONT_PATH = os.path.join(BASE_DIR, "Amiri.ttf")

def _execute_pdf_to_docx(pdf_path, docx_path):
    """تحويل PDF إلى Word بدقة عالية وبدون حظر البوت"""
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None, pages=None)
        cv.close()
        return True
    except Exception as e:
        logger.error(f"خطأ أثناء تحويل PDF إلى Word: {str(e)}")
        return False

async def convert_pdf_to_docx(pdf_path, docx_path):
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(executor, _execute_pdf_to_docx, pdf_path, docx_path)


def clean_and_reshape_arabic(text):
    """
    حل المشكلة الثالثة: التشكيل السياقي وعكس الاتجاه (RTL) للحروف العربية
    لضمان عدم ظهور النص كرموز مقلوبة أو مبعثرة
    """
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        
        # إعادة تشكيل الحروف لترتبط ببعضها سياقياً حسب موقعها في الكلمة
        reshaped_text = arabic_reshaper.reshape(text)
        # عكس اتجاه النص ليدعم القراءة الصحيحة من اليمين إلى اليسار
        bidi_text = get_display(reshaped_text)
        return bidi_text
    except ImportError:
        # حماية تراجعية في حال عدم توفر المكتبات بعد على السيرفر
        logger.warning("⚠️ حزم arabic_reshaper أو python-bidi غير مثبتة، سيتم إخراج النص الخام.")
        return text


def _execute_docx_to_pdf(docx_path, pdf_path):
    """
    معمارية تحويل داخلية مصححة تماماً (Fail-safe)
    تحل مشاكل: غياب الخط، دعم العربية، وتصحيح معاملات دالة PyMuPDF
    """
    try:
        import fitz  # PyMuPDF
        doc = Document(docx_path)
        pdf_doc = fitz.open()
        
        story = []
        
        # 1. قراءة النصوص والفقرات ومعالجتها عربياً فوراً
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                processed_text = clean_and_reshape_arabic(text)
                story.append(('paragraph', processed_text))
                
        # 2. قراءة الجداول ومعالجة نصوصها
        for table in doc.tables:
            for row in table.rows:
                cells_text = [clean_and_reshape_arabic(cell.text.strip()) for cell in row.cells if cell.text.strip()]
                if cells_text:
                    grid_line = " | ".join(cells_text)
                    story.append(('table_row', grid_line))

        # 3. معالجة الصور والأشكال المضمنة
        for inline_shape in doc.inline_shapes:
            if inline_shape.type in [WD_INLINE_SHAPE_TYPE.PICTURE, WD_INLINE_SHAPE_TYPE.LINKED_PICTURE]:
                img_msg = clean_and_reshape_arabic("[🖼️ صورة أو رسمة مضمنة في المستند الأصلي]")
                story.append(('image_placeholder', img_msg))

        if not story:
            story.append(('paragraph', clean_and_reshape_arabic("[مستند فارغ من النصوص]")))

        # التحقق من وجود الخط العربي المرفق بالمسار المطلق الصحيح
        has_font = os.path.exists(FONT_PATH)
        if not has_font:
            logger.critical(f"❌ خطأ فادح: ملف الخط العربي غير موجود في المسار المحدد: {FONT_PATH}")

        page_content = ""
        lines_on_page = 0
        max_lines_per_page = 22
        
        for element_type, content in story:
            if element_type == 'paragraph':
                page_content += f"{content}\n\n"
            elif element_type == 'table_row':
                page_content += f"📊 {content}\n\n"
            elif element_type == 'image_placeholder':
                page_content += f"{content}\n\n"
                
            lines_on_page += (len(content) // 55) + 2
            
            if lines_on_page >= max_lines_per_page:
                page = pdf_doc.new_page(width=595, height=842) # حجم A4
                
                if has_font:
                    # حقن الخط برمجياً بالهيكل الإنشائي للصفحة
                    page.insert_font(fontname="ArabicFont", fontfile=FONT_PATH)
                    page.insert_text((50, 50), text=page_content, fontsize=12, fontname="ArabicFont")
                else:
                    # تراجع آمن (Fall-back) بخط النظام الافتراضي في حال فقدان الملف تماماً
                    page.insert_text((50, 50), text=page_content, fontsize=11, fontname="helv")
                
                page_content = ""
                lines_on_page = 0

        # إنشاء الصفحة الأخيرة للمستند
        if page_content or len(pdf_doc) == 0:
            page = pdf_doc.new_page(width=595, height=842)
            if has_font:
                page.insert_font(fontname="ArabicFont", fontfile=FONT_PATH)
                final_text = page_content if page_content else clean_and_reshape_arabic("تم التحويل بنجاح")
                page.insert_text((50, 50), text=final_text, fontsize=12, fontname="ArabicFont")
            else:
                final_text = page_content if page_content else "Done"
                page.insert_text((50, 50), text=final_text, fontsize=11, fontname="helv")

        # حفظ وإغلاق التدفق
        pdf_doc.save(pdf_path)
        pdf_doc.close()
        return True

    except Exception as e:
        logger.error(f"خطأ داخلي حرج في معمارية التحويل المباشر: {str(e)}")
        return False

async def convert_docx_to_pdf(docx_path, pdf_path):
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(executor, _execute_docx_to_pdf, docx_path, pdf_path)
